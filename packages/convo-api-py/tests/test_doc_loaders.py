import tempfile
from typing import Dict, Optional

import boto3
import docx
import pytest
from convo.embeddings.embed_documents import (
    get_content_category,
    load_documents,
)
from convo.embeddings.s3_loader import load_s3
from convo.embeddings.types import DocumentEmbeddingRequest
from moto import mock_aws
from reportlab.pdfgen import canvas


@pytest.mark.parametrize("chunk_by_title", [True, False])
@pytest.mark.parametrize(
    "kwargs",
    [dict(), dict(content_type="application/vnd.openxmlformats-officedocument")],
)
@mock_aws
def test_s3_load(chunk_by_title: bool, kwargs: Dict):
    conn = boto3.resource("s3", region_name="us-east-1")
    conn.create_bucket(Bucket="bucket")

    title = "Title"
    body = "Foo bar."

    with tempfile.TemporaryDirectory() as tmpdir:
        document = docx.Document()
        document.add_paragraph(title, style="Heading 1")
        document.add_paragraph(body, style="Body Text")
        f_path = f"{tmpdir}/test.docx"
        document.save(f_path)

        s3 = boto3.client("s3", region_name="us-east-1")
        s3.upload_file(f_path, "bucket", "foo")

        s3_path = "s3://bucket/foo/test.docx"
        chunks = load_s3(s3_path, **kwargs)

        assert len(chunks) == 2
        assert chunks[0].text == title
        assert chunks[1].text == body

        content_type, _, content_category = get_content_category(
            DocumentEmbeddingRequest(
                location=f_path,
                contentType=kwargs.get("content_type"),
                chunk_by_title=chunk_by_title,
            ),
            chunks,
        )
        assert content_type == ".docx"
        content_category == "document"


@pytest.mark.parametrize("chunk_by_title", [True, False])
@pytest.mark.parametrize("content_type", [None, "text/markdown"])
def test_load_url(chunk_by_title: bool, content_type: Optional[str]):
    url = (
        "https://raw.githubusercontent.com/"
        "Unstructured-IO/unstructured/main/LICENSE.md"
    )

    request = DocumentEmbeddingRequest(
        location=url, contentType=content_type, chunk_by_title=chunk_by_title
    )

    chunks = load_documents(request)

    assert "Apache License" in chunks[0].text

    content_type, _, content_category = get_content_category(request, chunks)
    assert content_type == ".md"
    content_category == "document"


@pytest.mark.parametrize("chunk_by_title", [True, False])
def test_load_pdf_file(chunk_by_title: bool):
    with tempfile.TemporaryDirectory() as tmpdir:
        f_path = f"{tmpdir}/test.pdf"
        pdf = canvas.Canvas(f_path)
        title = "Title"
        text = "Hello World"
        pdf.setTitle(title)
        pdf.drawString(100, 100, text)
        pdf.save()

        request = DocumentEmbeddingRequest(
            location=f_path,
            contentType="application/pdf",
            chunk_by_title=chunk_by_title,
        )

        chunks = load_documents(request)

        assert len(chunks) == 1
        assert chunks[0].text == text

        content_type, _, content_category = get_content_category(request, chunks)
        assert content_type == ".pdf"
        content_category == "document"
